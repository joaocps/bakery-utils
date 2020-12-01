import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt

LOGGER = logging.getLogger(__name__)
LOGGER.setLevel(logging.DEBUG)


class WordDoc:
    def __init__(self, data, month):
        self.document = Document()
        self._data = data
        self._month = month

    def create(self):

        iterations = iter(self._data)

        table = self.document.add_table(rows=1, cols=2)
        table.width = '100%'

        for client in iterations:
            cells = table.add_row().cells
            try:
                """Left Column"""
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                left = cells[0].paragraphs[0].add_run()

                left.add_picture('padariacentral.png', width=1400000, height=1400000)
                client_text = cells[0].add_paragraph()
                client_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = client_text.add_run(client.nome)

                font_client_name = run.font
                font_client_name.size = Pt(15)
                font_client_name.bold = True

                client_month_value = client.month_verify(4)

                for month, value in client_month_value.items():
                    new_month = cells[0].add_paragraph()
                    new_month.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format = new_month.paragraph_format
                    paragraph_format.space_after = Pt(5)
                    new_month_data = new_month.add_run(month + ' - ' + str(value) + '€')

                if len(client_month_value) > 1:
                    client_month_total = cells[0].add_paragraph()
                    client_month_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    data_run = client_month_total.add_run(client.month_total(4))
                    underliner = data_run.font
                    underliner.underline = True

                if len(client_month_value) == 0:
                    right_client_month_total = cells[0].add_paragraph()
                    right_client_month_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    data_run_right = right_client_month_total.add_run("Todos os meses pagos até ao momento.")
                    underliner = data_run_right.font
                    underliner.underline = True

                client_footer = cells[0].add_paragraph()
                client_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = client_footer.add_run("Padaria Central - Sangalhos - 234741357/969015589")

                font_footer = run.font
                font_footer.size = Pt(7)
                font_footer.italic = True
                styler = client_footer.paragraph_format
                styler.space_after = Pt(5)

                """ Right Column"""
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                right = cells[1].paragraphs[0].add_run()

                right_client = next(iterations)
                right.add_picture('padariacentral.png', width=1400000, height=1400000)
                right_client_text = cells[1].add_paragraph()
                right_client_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = right_client_text.add_run(right_client.nome)

                font_client_name = run.font
                font_client_name.size = Pt(15)
                font_client_name.bold = True

                right_client_month_value = right_client.month_verify(4)

                for month, value in right_client_month_value.items():
                    new_month = cells[1].add_paragraph()
                    new_month.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format = new_month.paragraph_format
                    paragraph_format.space_after = Pt(5)
                    new_month_data = new_month.add_run(month + ' - ' + str(value) + '€')

                if len(right_client_month_value) > 1:
                    right_client_month_total = cells[1].add_paragraph()
                    right_client_month_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    data_run_right = right_client_month_total.add_run(right_client.month_total(4))
                    underliner = data_run_right.font
                    underliner.underline = True

                if len(right_client_month_value) == 0:
                    right_client_month_total = cells[1].add_paragraph()
                    right_client_month_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    data_run_right = right_client_month_total.add_run("Todos os meses pagos até ao momento.")
                    underliner = data_run_right.font
                    underliner.underline = True

                right_client_footer = cells[1].add_paragraph()
                right_client_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = right_client_footer.add_run("Padaria Central - Sangalhos - 234741357/969015589")

                font_footer = run.font
                font_footer.size = Pt(7)
                font_footer.italic = True

                styler = right_client_footer.paragraph_format
                styler.space_after = Pt(5)


            except StopIteration:
                right_client = None
                # End space introduction

        self.document.save('demo.docx')
