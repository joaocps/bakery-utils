import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE

LOGGER = logging.getLogger(__name__)
LOGGER.setLevel(logging.DEBUG)


class Word3Cols:
    def __init__(self, data, month, output_path, op, info):
        self.document = Document()
        self._data = data
        self._month = month
        self._output_path = output_path
        self._op = op
        self._info = info

    def create(self):

        """Process data and remove some clients"""
        desired_clients = []
        for x in self._data:
            _ = x.month_verify(self._month)
            if len(_) != 0 and self._op == "pay":
                desired_clients.append(x)

        """Iterate over desire clients"""
        iterations = iter(desired_clients)

        """changing the page margins"""
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(0)
            section.bottom_margin = Cm(0)
            section.left_margin = Cm(0)
            section.right_margin = Cm(0)

        """Table creation"""
        table = self.document.add_table(rows=0, cols=3)
        table.width = '100%'

        """Creation Main"""
        for client in iterations:
            cells = table.add_row().cells
            try:
                """Left column"""
                self.add_image(cells=cells, col=0)
                self.add_paragraph(content=client.nome, cells=cells, col=0,
                                   font_size=15, set_bold=True, space_after=5)

                if self._op == 'pay':
                    l_months = client.month_verify(self._month)

                    if len(l_months) <= 2:
                        for month, value in l_months.items():
                            self.add_paragraph(month + ' - ' + str(value) + '€', cells=cells, col=0, font_size=12)
                    else:
                        print(l_months)

                    if len(l_months) == 0:
                        self.add_paragraph("Todos os meses pagos até ao momento.", cells=cells, col=0,
                                           set_underline=True, font_size=12)

                    if len(l_months) >= 2:
                        self.add_paragraph(client.month_total(self._month), cells=cells, col=0, set_underline=True, font_size=12,
                                           space_before=5)
                else:
                    self.add_paragraph(self._info, cells=cells, col=0, set_underline=True, font_size=10)

                """CellFooter"""
                self.add_paragraph("Padaria Central - Sangalhos - 234741357/969015589",
                                   cells=cells, col=0, set_italic=True, font_size=7, space_before=5, space_after=10)

                """Mid column"""
                mid_client = next(iterations)
                self.add_image(cells=cells, col=1)
                self.add_paragraph(content=mid_client.nome, cells=cells, col=1,
                                   font_size=15, set_bold=True, space_after=5)

                if self._op == 'pay':
                    m_months = mid_client.month_verify(self._month)

                    if len(m_months) <= 2:
                        for month, value in m_months.items():
                            self.add_paragraph(month + ' - ' + str(value) + '€', cells=cells, col=1, font_size=12)
                    else:
                        print(m_months)

                    if len(m_months) == 0:
                        self.add_paragraph("Todos os meses pagos até ao momento.", cells=cells, col=1,
                                           set_underline=True, font_size=12)

                    if len(m_months) >= 2:
                        self.add_paragraph(mid_client.month_total(self._month), cells=cells, col=1, set_underline=True, font_size=12,
                                           space_before=5)
                else:
                    self.add_paragraph(self._info, cells=cells, col=1, set_underline=True, font_size=10)

                """CellFooter"""
                self.add_paragraph("Padaria Central - Sangalhos - 234741357/969015589",
                                   cells=cells, col=1, set_italic=True, font_size=7, space_before=5, space_after=10)

                """Right column"""
                right_client = next(iterations)
                self.add_image(cells=cells, col=2)
                self.add_paragraph(content=right_client.nome, cells=cells, col=2,
                                   font_size=15, set_bold=True, space_after=5)

                if self._op == 'pay':
                    r_months = right_client.month_verify(self._month)

                    if len(r_months) <= 2:
                        for month, value in r_months.items():
                            self.add_paragraph(month + ' - ' + str(value) + '€', cells=cells, col=2, font_size=12)
                    else:
                        print(r_months)

                    if len(r_months) == 0:
                        self.add_paragraph("Todos os meses pagos até ao momento.", cells=cells, col=2,
                                           set_underline=True, font_size=12)

                    if len(r_months) >= 2:
                        self.add_paragraph(right_client.month_total(self._month), cells=cells, col=2, set_underline=True, font_size=12,
                                           space_before=5)
                else:
                    self.add_paragraph(self._info, cells=cells, col=2, set_underline=True, font_size=10)

                """CellFooter"""
                self.add_paragraph("Padaria Central - Sangalhos - 234741357/969015589",
                                   cells=cells, col=2, set_italic=True, font_size=7, space_before=5, space_after=10)

            except StopIteration:
                mid_client = None
                right_client = None
                # Stupid but works

        self.document.save(self._output_path)

    """Utils to add paragraph"""
    def add_paragraph(self, content, cells, col, space_after=0, font_name='Arial', font_size=16, line_spacing=0,
                      space_before=0, keep_together=True, keep_with_next=False, page_break_before=False,
                      widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):

        paragraph = cells[col].add_paragraph()
        paragraph_content = paragraph.add_run(content)
        font = paragraph_content.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control


    def add_image(self, cells, col):
        cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[col].paragraphs[0].add_run().add_picture('padariacentral.png', width=1400000, height=1400000)
